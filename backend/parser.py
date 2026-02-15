"""
Document Parser Module
======================
Handles extraction of text content from PDF and DOCX files.
Supports resume parsing with error handling and text cleaning.

Author: AI Resume Matcher Team
Version: 1.0.0
"""

import io
import re
import logging
from pathlib import Path
from typing import Optional, Union

import pdfplumber
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentParser:
    """
    A class to parse PDF and DOCX documents and extract clean text.
    
    Attributes:
        supported_formats (list): List of supported file extensions
    """
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc']
    
    def __init__(self):
        """Initialize the DocumentParser with logging configuration."""
        self.logger = logging.getLogger(self.__class__.__name__)
        self.logger.info("DocumentParser initialized")
    
    def parse_file(self, file_path: Union[str, Path]) -> dict:
        """
        Parse a document file and extract text content.
        
        Args:
            file_path: Path to the document file (PDF or DOCX)
            
        Returns:
            dict: Contains 'text' (extracted text), 'filename', 'format', 
                  'page_count' (for PDFs), and 'success' status
                  
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        file_path = Path(file_path)
        
        # Validate file existence
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file format
        file_extension = file_path.suffix.lower()
        if file_extension not in self.SUPPORTED_FORMATS:
            self.logger.error(f"Unsupported file format: {file_extension}")
            raise ValueError(
                f"Unsupported file format: {file_extension}. "
                f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            )
        
        self.logger.info(f"Parsing file: {file_path.name}")
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(file_path)
        except Exception as e:
            self.logger.error(f"Error parsing file {file_path.name}: {str(e)}")
            return {
                'text': '',
                'filename': file_path.name,
                'format': file_extension,
                'page_count': 0,
                'success': False,
                'error': str(e)
            }
    
    def parse_bytes(self, file_bytes: bytes, filename: str) -> dict:
        """
        Parse document from bytes (for uploaded files).
        
        Args:
            file_bytes: Raw bytes of the document
            filename: Name of the file (used to determine format)
            
        Returns:
            dict: Contains extracted text and metadata
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        self.logger.info(f"Parsing file from bytes: {filename}")
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf_bytes(file_bytes, filename)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx_bytes(file_bytes, filename)
        except Exception as e:
            self.logger.error(f"Error parsing {filename}: {str(e)}")
            return {
                'text': '',
                'filename': filename,
                'format': file_extension,
                'page_count': 0,
                'success': False,
                'error': str(e)
            }
    
    def _parse_pdf(self, file_path: Path) -> dict:
        """
        Extract text from PDF file using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            dict: Extracted text and metadata
        """
        text_content = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            self.logger.info(f"PDF has {page_count} pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        self.logger.debug(f"Extracted text from page {page_num}")
                except Exception as e:
                    self.logger.warning(f"Error extracting page {page_num}: {e}")
        
        raw_text = '\n'.join(text_content)
        cleaned_text = self._clean_text(raw_text)
        
        self.logger.info(f"Successfully extracted {len(cleaned_text)} characters from PDF")
        
        return {
            'text': cleaned_text,
            'filename': file_path.name,
            'format': '.pdf',
            'page_count': page_count,
            'success': True,
            'error': None
        }
    
    def _parse_pdf_bytes(self, file_bytes: bytes, filename: str) -> dict:
        """
        Extract text from PDF bytes.
        
        Args:
            file_bytes: Raw PDF bytes
            filename: Original filename
            
        Returns:
            dict: Extracted text and metadata
        """
        text_content = []
        page_count = 0
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            self.logger.info(f"PDF has {page_count} pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    self.logger.warning(f"Error extracting page {page_num}: {e}")
        
        raw_text = '\n'.join(text_content)
        cleaned_text = self._clean_text(raw_text)
        
        return {
            'text': cleaned_text,
            'filename': filename,
            'format': '.pdf',
            'page_count': page_count,
            'success': True,
            'error': None
        }
    
    def _parse_docx(self, file_path: Path) -> dict:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            dict: Extracted text and metadata
        """
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text_content = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(' '.join(row_text))
        
        raw_text = '\n'.join(text_content)
        cleaned_text = self._clean_text(raw_text)
        
        self.logger.info(f"Successfully extracted {len(cleaned_text)} characters from DOCX")
        
        return {
            'text': cleaned_text,
            'filename': file_path.name,
            'format': '.docx',
            'page_count': len(doc.sections),
            'success': True,
            'error': None
        }
    
    def _parse_docx_bytes(self, file_bytes: bytes, filename: str) -> dict:
        """
        Extract text from DOCX bytes.
        
        Args:
            file_bytes: Raw DOCX bytes
            filename: Original filename
            
        Returns:
            dict: Extracted text and metadata
        """
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract text from all paragraphs
        text_content = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(' '.join(row_text))
        
        raw_text = '\n'.join(text_content)
        cleaned_text = self._clean_text(raw_text)
        
        return {
            'text': cleaned_text,
            'filename': filename,
            'format': '.docx',
            'page_count': len(doc.sections),
            'success': True,
            'error': None
        }
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing special characters and normalizing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ''
        
        # Remove extra whitespace and normalize
        cleaned = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        cleaned = re.sub(r'[^\w\s.,;:!?@\-()/+]', ' ', cleaned)
        
        # Remove extra spaces
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        # Strip leading/trailing whitespace
        cleaned = cleaned.strip()
        
        return cleaned
    
    def extract_sections(self, text: str) -> dict:
        """
        Attempt to extract common resume sections from text.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            dict: Dictionary with section names as keys
        """
        sections = {
            'full_text': text,
            'contact_info': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'certifications': ''
        }
        
        # Common section headers
        section_patterns = {
            'contact_info': r'(?i)(contact|personal info|profile)',
            'summary': r'(?i)(summary|objective|about|profile)',
            'experience': r'(?i)(experience|work|employment|career|professional)',
            'education': r'(?i)(education|academic|qualification|degree)',
            'skills': r'(?i)(skills|technologies|competencies|expertise)',
            'certifications': r'(?i)(certifications|certificates|credentials)'
        }
        
        # Simple section extraction based on headers
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if line is a section header
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_stripped, re.IGNORECASE):
                    current_section = section_name
                    break
            
            # Add line to current section
            if current_section and line_stripped:
                sections[current_section] += line_stripped + ' '
        
        return sections


# Convenience function for quick parsing
def parse_resume(file_path: Union[str, Path]) -> dict:
    """
    Quick function to parse a resume file.
    
    Args:
        file_path: Path to resume file
        
    Returns:
        dict: Parsed content and metadata
    """
    parser = DocumentParser()
    return parser.parse_file(file_path)


def parse_resume_bytes(file_bytes: bytes, filename: str) -> dict:
    """
    Quick function to parse resume from bytes.
    
    Args:
        file_bytes: Raw file bytes
        filename: Original filename
        
    Returns:
        dict: Parsed content and metadata
    """
    parser = DocumentParser()
    return parser.parse_bytes(file_bytes, filename)


if __name__ == "__main__":
    # Test the parser
    import sys
    
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        result = parse_resume(test_file)
        
        print(f"\n{'='*60}")
        print(f"Parsing Result for: {result['filename']}")
        print(f"{'='*60}")
        print(f"Format: {result['format']}")
        print(f"Pages/Sections: {result['page_count']}")
        print(f"Success: {result['success']}")
        print(f"Text Length: {len(result['text'])} characters")
        print(f"\n--- First 500 characters ---")
        print(result['text'][:500])
        print(f"\n{'='*60}")
    else:
        print("Usage: python parser.py <path_to_resume_file>")
